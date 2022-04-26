# -*- coding: utf-8 -*-
"""
Created on Sun Feb  4 20:36:57 2018

@author: user
"""

import tensorflow as tf
import numpy
import csv
import matplotlib.pyplot as plt
import pandas as pd
import smtplib
import email
import email.mime.application
import email.mime.multipart
from sklearn.preprocessing import OneHotEncoder
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from docx import Document

example_arr = []
example_arr_X = []
example_arr_y = []
accuracy_arr = []

per_epoch_batches = 68

curr_batch = 1   

data_check_count = 1 

n_nodes_hl1 = 12
n_nodes_hl2 = 12
n_nodes_hl3 = 12

n_classes = 2

x = tf.placeholder('float',[None,23])
y = tf.placeholder('float')

#This function is used to read the dataset and store in a variable called as features 
#which will be used ahead. For dataset reading in tensorflow a filename queue is required 
#unlike the conventional pandas.read_csv method
def read_from_dataset():
    
    global features
    
    filename_queue = tf.train.string_input_producer(["C:/xampp5/htdocs/be_project/dataset/test2.csv"])
    
    reader = tf.TextLineReader(skip_header_lines=1)
    key, value = reader.read(filename_queue)
    
    # Default values, in case of empty columns. Also specifies the type of the
    # decoded result.
    record_defaults = [[1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1], [1]]
    col1, col2, col3, col4, col5, col6,col7, col8, col9, col10, col11, col12, col13, col14, col15, col16, col17, col18, col19, col20, col21, col22, col23, col24 = tf.decode_csv(
        value, record_defaults=record_defaults)
    features = tf.stack([col1, col2, col3, col4, col5, col6,col7, col8, col9, col10, col11, col12, col13, col14, col15, col16, col17, col18, col19, col20, col21, col22, col23, col24])   

#This function is used to iterate and get the values of the next batch ones the current
#batch is processed    
def iterate_through_batch():
    
    with tf.Session() as sess:
        global curr_batch
        
        for _ in range(curr_batch):
            
            sess.run(next_element_x)
            sess.run(next_element_y)

        batch_x = next_element_x.eval()
        batch_y = next_element_y.eval()
        
        curr_batch += 1
        
        return batch_x,batch_y
    
#Backpropagation requires that the data should be provided in batches and hence this function
#creates batches from the entire dataset
def create_batch():
    
    global next_element_x, next_element_y
    
    X_train_new = tf.data.Dataset.from_tensor_slices(X_train)
    y_train_new = tf.data.Dataset.from_tensor_slices(y_train)
  
    batch_x = X_train_new.batch(100)
    batch_y = y_train_new.batch(100)
    
    iterator_x = batch_x.make_one_shot_iterator()
    next_element_x = iterator_x.get_next()
    iterator_y = batch_y.make_one_shot_iterator()
    next_element_y = iterator_y.get_next()
    
#This a method which is used at the initial stage before batch creation. It does the work
#of preparing the dataset in program for creation of batches further ahead
def prepare_dataset():
    
    with tf.Session() as sess:
        
        # Start populating the filename queue.
        
        global dataset, X_test, y_test, X_train, y_train
        coord = tf.train.Coordinator()
        threads = tf.train.start_queue_runners(coord=coord)
    
        for i in range(8519):
            
            example_arr_X.append(sess.run(features))
            
        coord.request_stop()
        coord.join(threads)
      
        dataset = numpy.asarray(example_arr_X)
        
        onehotencoder = OneHotEncoder(categorical_features = [23])
        dataset = onehotencoder.fit_transform(dataset).toarray()
        
        temp_dataset = numpy.hsplit(dataset, [2,25])
        X_total = temp_dataset[1]
        y_total = temp_dataset[0]
        
        X_train, X_test, y_train, y_test = train_test_split(X_total, y_total, test_size = 0.2, random_state = 42)
        
        sc = StandardScaler()
        X_train = sc.fit_transform(X_train)
        X_test = sc.transform(X_test)

#This function creates the structure of the neural network
def neural_net_model(data):
    
    global output_layer, hidden_1_layer, hidden_2_layer, hidden_3_layer
    hidden_1_layer = {'weights':tf.Variable(tf.random_normal([23,n_nodes_hl1])),'biases':tf.Variable(tf.random_normal([n_nodes_hl1]))}
    hidden_2_layer = {'weights':tf.Variable(tf.random_normal([n_nodes_hl1,n_nodes_hl2])),
		'biases':tf.Variable(tf.random_normal([n_nodes_hl2]))}
    hidden_3_layer = {'weights':tf.Variable(tf.random_normal([n_nodes_hl2,n_nodes_hl3])),
		'biases':tf.Variable(tf.random_normal([n_nodes_hl3]))}
    output_layer = {'weights':tf.Variable(tf.random_normal([n_nodes_hl3,n_classes])),
		'biases':tf.Variable(tf.random_normal([n_classes]))}

    l1 = tf.add(tf.matmul(data, hidden_1_layer['weights']),hidden_1_layer['biases'])
    l1 = tf.nn.relu(l1)

    l2 = tf.add(tf.matmul(l1, hidden_2_layer['weights']),hidden_2_layer['biases'])
    l2 = tf.nn.relu(l2)
	
    l3 = tf.add(tf.matmul(l2, hidden_3_layer['weights']),hidden_3_layer['biases'])
    l3 = tf.nn.sigmoid(l3)

    output = tf.matmul(l3, output_layer['weights']) + output_layer['biases']

    return output

#This function makes use of the structure of neural network to pass values through it and
#update weights of the layers. Then finallyit calculates the probability of the new data    
def train_neural_network(x):
    
    global X_test,y_test, curr_batch,accuracy_arr, y_pred, output_probability, probs, output_single, prob_single, output_layer, hidden_1_layer, hidden_2_layer, hidden_3_layer
    prediction = neural_net_model(x)
    cost = tf.reduce_mean( tf.nn.softmax_cross_entropy_with_logits(labels = y, logits = prediction))
    output_probability = tf.nn.softmax(prediction)
    optimizer = tf.train.RMSPropOptimizer(learning_rate=0.001).minimize(cost)

    hm_epochs = 25
                            
    with tf.Session() as sess:
        
        sess.run(tf.global_variables_initializer())

        for epoch in range(hm_epochs):
            
            epoch_loss = 0
            iteration = 0
            create_batch()
            
            for _ in range(per_epoch_batches):
                
                epoch_x,epoch_y  = iterate_through_batch()
                _,epoch_c = sess.run([optimizer, cost], feed_dict = {x: epoch_x, y: epoch_y})
                epoch_loss += epoch_c
                print("Total loss at ", iteration+1," iteration : ",epoch_loss)
                iteration = iteration+1
                
            print('Epoch', epoch+1, 'completed out of ', hm_epochs, 'loss: ', epoch_loss)
            correct = tf.equal(tf.argmax(prediction, 1), tf.argmax(y,1))
            accuracy = tf.reduce_mean(tf.cast(correct, 'float'))*100
            print('Accuracy:', accuracy.eval({x: X_test, y: y_test}))
            curr_batch = 1
            accuracy_arr.append(accuracy.eval({x: X_test, y: y_test}))
            output = sess.run(tf.argmax(prediction, 1), {x:X_test})

            probs = sess.run(output_probability, feed_dict={x: X_test})*100
            y_pred = output
    
        new_dataset = pd.read_csv("C:/xampp5/htdocs/be_project/dataset/temp.csv", header = None)        
        X_new = new_dataset.iloc[:, :].values
        sc = StandardScaler()
        X_new = sc.fit_transform(X_new)
        output_single = sess.run(tf.argmax(prediction, 1), {x:X_new})
        prob_single = sess.run(output_probability, feed_dict={x: X_new})*100
        print(output_single)
        print(prob_single)
        for i in range(len(prob_single)):  
            prob_single_send = prob_single[i]
            prob_single_send = prob_single_send[1]
            write_to_docx(prob_single_send)
            send_mail()
        clear_temp()
        quit()
    
#Used for writing the result into a docx file along with other data analysis images
def write_to_docx(prob):
    
    document = Document()
    intro_line = 'Thank you for using our predictor!'
    document.add_paragraph(intro_line)
    percentage_line = 'Your percentage of risk of having cancer is '+str(prob)+'%'
    document.add_paragraph(percentage_line)
    diagram_intro = 'We also have done a detailed analysis of how lung cancer is distributed among the people to gain a deeper insight. Please have a look at it too'
    document.add_paragraph(diagram_intro)
    document.add_picture('histogram.png')
    document.add_picture('pie_chart.png')
    document.save('C:/xampp5/htdocs/be_project/python/demo.docx')
    print("written")
    
#clear the file in which new data is stored in order to allow newer data 
def clear_temp():
    
    myFile = open('C:/xampp5/htdocs/be_project/dataset/temp.csv', 'w')
    email_file = open('C:/xampp5/htdocs/be_project/dataset/emails.csv', 'w')
    
    with myFile:
        writer = csv.writer(myFile)
        writer.writerows("")
        print("Clearing complete")
        
    with email_file:
        writer = csv.writer(email_file)
        writer.writerows("")
        print("Clearing complete")
          
#A function used for visualizing the accuracy curve. Accuracies after every epoch are stored
def visualize_accuracy():
    
    global accuracy_graph
    epoch = []
    accuracy = []
    
    fig = plt.figure()
    ax = fig.add_subplot(111)
    
    for index in range(len(accuracy_arr)):
        
        epoch.append(index)
        accuracy.append(accuracy_arr[index])
        
    ax.set_title('Accuracy change with respect to epochs')
    ax.set_xlabel('Accuracy(%)')
    ax.set_ylabel('Epochs')
    ax.plot(accuracy,epoch)
    
    return
    
#Used to check if the file containing new data is empty or any new data is arrived and
#waiting to be processed
def check_new_data():
    
    try:
        pd.read_csv("C:/xampp5/htdocs/be_project/dataset/temp.csv")
    except:
        return 0
    
    return 1

#Used to send the mail of the result docx file to the end user
def send_mail():
    email_dataset = pd.read_csv("C:/xampp5/htdocs/be_project/dataset/emails.csv", header = None)
    user_email = email_dataset.iloc[:, :].values
    for i in range(len(user_email)): 
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login("anvaykhopkar1234@gmail.com", "neilgupte161996")
        
        msg = email.mime.multipart.MIMEMultipart()
        filename='demo.docx'
        fp=open(filename,'rb')
        att = email.mime.application.MIMEApplication(fp.read(),_subtype="docx")
        fp.close()
        att.add_header('Content-Disposition','attachment',filename=filename)
        msg.attach(att)
    
        server.sendmail("sumitkhopkar25@gmail.com", user_email[i], msg.as_string())
        server.quit()
    
def run_ann():   
    read_from_dataset()
    prepare_dataset()
    train_neural_network(x)
    
#Entry point of program
def main():
    
    global data_check_count
    
    while(data_check_count < 500):
        
        if(check_new_data() != 1):
            
            print ("No new data")
            data_check_count += 1
            main()
        
        else:
            run_ann()
    
    quit()  
        
if __name__== "__main__":
    main()